import streamlit as st
import requests
import pandas as pd
import datetime
import moviepy.editor as mp
from PIL import Image
import docx
from results import *
import io
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from moviepy.video.io.VideoFileClip import VideoFileClip
import threading
import os
import numpy as np


# Set the title of the app
st.title("Video Meeting Summarizer")

if 'start_point' not in st.session_state:
    st.session_state['start_point'] = 0

def update_start(start_t):
    st.session_state['start_point'] = int(start_t / 1000)

uploaded_file = st.file_uploader("choose file", type=["webm" , "mp4"])
if uploaded_file is not None:
    with open("temp_file.webm", "wb") as f:
        f.write(uploaded_file.getbuffer())

    clip = VideoFileClip("temp_file.webm")
    duration_sec = int(clip.duration)

    hours, remainder = divmod(duration_sec, 3600)
    minutes, seconds = divmod(remainder, 60)

    # play video file
    st.video(uploaded_file, start_time=0)

    # get recipient email
    Agenda = st.text_input('Enter agenda of the meet')

    # get number of attendees
    num_attendees = st.number_input('Number of attendees', min_value=1, value=1)

    # get names of speakers
    speaker_names = st.text_input('Enter the names of the speakers (comma-separated)')

    # get recipient email
    recipient_email = recipient_email = st.text_input('Enter your email address')
    
    def generate_summary():
        # upload video to AssemblyAI
        polling_endpoint = upload_to_AssemblyAI(uploaded_file)

        status = 'submitted'
        while status != "completed":
            polling_response = requests.get(polling_endpoint, headers=headers)
            status = polling_response.json()['status']

            if status == 'completed':

                # display chapter summaries
                st.subheader('')
                chapters = polling_response.json()['chapters']
                chapters_df = pd.DataFrame(chapters)

                # create text file with chapter summaries

                summary_text = f"\n\nDate:{datetime.datetime.now().strftime('%d-%m-%Y')}\n"
                summary_text += f"Number of attendees: {num_attendees}\n"
                summary_text += "Duration: {:02d}:{:02d}:{:02d}\n".format(hours, minutes, seconds)
                summary_text += f"Names of the speaker: {speaker_names}\n\n"
                summary_text += f"Agenda: {Agenda}\n\n"
                summary_text += f"Points discussed in the meet:\n\n"

                # take snapshots randomly from the video
                snapshot_times = [int(duration_sec * r) for r in np.random.uniform(size=1)]
                snapshots = []
                for t in snapshot_times:
                    try:
                        snapshot = clip.save_frame(f"snapshot_{t}.png", t=t)
                        snapshots.append(snapshot)
                    except:
                        pass

                for index, row in chapters_df.iterrows():
                    summary_text += f"-- {row['gist']}\n"
                    summary_text += f"   {row['summary']}\n\n"

                doc = docx.Document()
                heading = doc.add_heading('Minutes of the meet', level=1)
                heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(summary_text)
                images = [file for file in os.listdir() if file.endswith(f'_{t}.png')]
                doc.add_paragraph('Snapshot images:')
                for image_path in images:
                    try:
                        img = Image.open(image_path)
                        doc.add_picture(image_path, width=docx.shared.Inches(6))
                    except:
                        print(f"Failed to add image {image_path}")

                doc_bytes = io.BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)


                summary_gist = "Meeting summary points:\n"
                for index, row in chapters_df.iterrows():
                    summary_gist += f"-- {row['gist']}\n"

                # Replace the following variables with appropriate values
                sender_email = "ashiqashervegar.2703@gmail.com"
                sender_password = "wcgfigqopnfchfmw"
                receiver_email = recipient_email
                subject = "Summary Generated"
                body = "Hello, Your summary is ready!!!!! \n\n" + summary_gist + "\n\nPlease find the attachment to go through the full summary"

                # Attach the previously generated .doc file

                attachment = MIMEApplication(doc_bytes.read(), _subtype='doc')
                attachment.add_header('Content-Disposition', 'attachment', filename="meeting_summary.doc")

                # Create the email message
                message = MIMEMultipart()
                message['From'] = sender_email
                message['To'] = receiver_email
                message['Subject'] = subject

                message.attach(MIMEText(body))
                message.attach(attachment)

                # Send the email using the SMTP server of your email provider
                with smtplib.SMTP('smtp.gmail.com', 587) as smtp:
                    smtp.ehlo()
                    smtp.starttls()
                    smtp.ehlo()
                    smtp.login(sender_email, sender_password)
                    smtp.sendmail(sender_email, receiver_email, message.as_string())


    if st.button("Generate Summary"):
        st.success('Summary document will be sent to your email shortly.')
        thread = threading.Thread(target=generate_summary)
        thread.start()
